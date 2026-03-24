"""
Generador de Informes de Madurez Digital - Web App
===================================================
Aplicación web para generar automáticamente los informes de diagnóstico
de madurez digital a partir del Excel backend y la plantilla Word.

Para ejecutar localmente:
    pip install streamlit openpyxl python-docx matplotlib numpy
    streamlit run app.py

Para desplegar en Streamlit Cloud:
    1. Sube este proyecto a GitHub
    2. Ve a share.streamlit.io
    3. Conecta tu repositorio
"""
import streamlit as st
import os, shutil, re, zipfile, tempfile, subprocess
from datetime import datetime
from io import BytesIO
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

# ─── Page config ───
st.set_page_config(
    page_title="Generador Informes Madurez Digital",
    page_icon="📊",
    layout="wide"
)

# ─── Constants ───
BLOCK_NAMES = [
    'Cultura Organizativa',
    'Identificación y Digitalización de Procesos',
    'Infraestructura Digital y Tecnológica',
    'Comunicación Digital y Marketing',
    'Seguridad Informática y Protección de Datos',
    'Capacitación Digital del Personal',
    'Innovación',
]
BLOCK_SHORT = ['Capaci-\ntación', 'Innovación', 'Cultura\nOrg.', 'Procesos',
               'Infraes-\ntructura', 'Comuni-\ncación', 'Seguridad']
MAX_COL = 188  # Column GF - end of Datos_para_Inf_word range


# ─── Helper functions ───

def safe_float(val, default=0):
    if val is None:
        return default
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip()
    if s.startswith('#') or s == '' or s == 'None':
        return default
    try:
        return float(s)
    except (ValueError, TypeError):
        return default


def get_entities(backend_bytes):
    """Read entity names from uploaded Excel."""
    wb = openpyxl.load_workbook(BytesIO(backend_bytes), data_only=True)
    ws = wb['Form1']
    entities = []
    for row in range(5, ws.max_row + 1):
        name = ws.cell(row=row, column=9).value
        if name and str(name).strip():
            entities.append(str(name))
    wb.close()
    return entities


def recalc_for_entity(backend_bytes, entity_name, work_dir):
    """Set entity in Diagnostico and recalculate."""
    work_file = os.path.join(work_dir, "backend_work.xlsx")
    with open(work_file, 'wb') as f:
        f.write(backend_bytes)
    wb = openpyxl.load_workbook(work_file)
    ws = wb['Diagnostico']
    ws['A9'] = entity_name
    wb.save(work_file)
    wb.close()
    # Try LibreOffice recalc
    try:
        subprocess.run(
            ['libreoffice', '--headless', '--calc', '--convert-to', 'xlsx',
             '--outdir', work_dir, work_file],
            capture_output=True, timeout=60
        )
    except Exception:
        pass  # If LibreOffice not available, use cached values
    return work_file


def read_all_fields(work_file):
    wb = openpyxl.load_workbook(work_file, data_only=True)
    ws = wb['Diagnostico']
    fields = {}
    for col in range(1, MAX_COL + 1):
        label = ws.cell(row=8, column=col).value
        value = ws.cell(row=9, column=col).value
        if label and isinstance(label, str) and label.strip():
            merge_name = label.strip().replace(' ', '_')
            if value is not None:
                if isinstance(value, datetime):
                    fields[merge_name] = value.strftime('%d/%m/%Y')
                else:
                    fields[merge_name] = str(value)
            else:
                fields[merge_name] = ''
    wb.close()
    return fields


def read_chart_data(work_file):
    wb = openpyxl.load_workbook(work_file, data_only=True)
    ws = wb['Gráficos 2.0']
    blocks = []
    for i in range(7):
        base_row = 3 + i * 4
        labels, max_scores, entity_scores = [], [], []
        for col in range(3, 9):
            lbl = ws.cell(row=base_row, column=col).value
            mx = ws.cell(row=base_row + 1, column=col).value
            sc = ws.cell(row=base_row + 2, column=col).value
            if lbl:
                labels.append(str(lbl))
                max_scores.append(safe_float(mx, 50))
                entity_scores.append(safe_float(sc, 0))
        blocks.append({
            'name': BLOCK_NAMES[i],
            'labels': labels,
            'max_scores': max_scores,
            'entity_scores': entity_scores,
            'total': safe_float(ws.cell(row=base_row + 2, column=9).value, 0),
            'max': safe_float(ws.cell(row=base_row, column=10).value, 300),
        })
    summary_pcts = []
    for col in range(2, 9):
        pct = ws.cell(row=33, column=col).value
        summary_pcts.append(safe_float(pct, 0))
    wb.close()
    return blocks, summary_pcts


def generate_block_chart(block, entity_name, output_path):
    plt.rcParams['font.family'] = 'DejaVu Sans'
    labels = block['labels']
    n = len(labels)
    angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
    angles += angles[:1]
    max_vals = block['max_scores'] + [block['max_scores'][0]]
    ent_vals = block['entity_scores'] + [block['entity_scores'][0]]

    fig, ax = plt.subplots(figsize=(4.2, 4.8), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('white')
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)

    # Hexagonal grid
    ax.yaxis.grid(False)
    ax.xaxis.grid(False)
    ax.spines['polar'].set_visible(False)
    max_val = max(max(max_vals), 50)
    for level in np.linspace(0, max_val, 6)[1:]:
        ax.plot(angles, [level] * (n + 1), color='#CCCCCC', linewidth=0.5, zorder=0)
    for angle in angles[:-1]:
        ax.plot([angle, angle], [0, max_val + 5], color='#CCCCCC', linewidth=0.4, zorder=0)

    ax.plot(angles, max_vals, color='#555555', linewidth=2.2, zorder=2,
            label='Rango de puntuación Máximo Posible')
    ax.plot(angles, ent_vals, color='#C2185B', linewidth=2.2, zorder=3,
            label=entity_name[:50])

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=8.5, color='#444444')
    ax.set_ylim(0, max_val + 5)
    ax.set_yticklabels([])
    ax.legend(fontsize=6.5, loc='lower center', bbox_to_anchor=(0.5, 1.05),
              ncol=1, frameon=False)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()


def generate_summary_chart(summary_pcts, entity_name, output_path):
    plt.rcParams['font.family'] = 'DejaVu Sans'
    labels = BLOCK_SHORT
    n = len(labels)
    angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
    angles += angles[:1]
    values = [p * 100 for p in summary_pcts] + [summary_pcts[0] * 100]
    max_values = [100] * (n + 1)

    fig, ax = plt.subplots(figsize=(4.5, 4.5), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('white')
    ax.set_theta_offset(0)
    ax.set_theta_direction(-1)
    ax.plot(angles, max_values, color='#555555', linewidth=1.8)
    ax.plot(angles, values, color='#C2185B', linewidth=2.0)
    ax.fill(angles, values, alpha=0.08, color='#C2185B')
    for angle, val in zip(angles[:-1], values[:-1]):
        ax.text(angle, val + 8, f'{val:.0f}%', ha='center', va='bottom',
                fontsize=8, color='#C2185B', fontweight='bold')
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=7.5, color='#333333')
    ax.set_ylim(0, 115)
    ax.set_yticklabels([])
    ax.xaxis.grid(True, color='#BBBBBB', linewidth=0.4)
    ax.yaxis.grid(True, color='#CCCCCC', linewidth=0.3)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()


def fill_merge_fields(doc, field_map):
    def process_element(element):
        all_elems = list(element.iter())
        field_starts = []
        for i, elem in enumerate(all_elems):
            if elem.tag == qn('w:fldChar'):
                if elem.get(qn('w:fldCharType')) == 'begin':
                    field_starts.append(i)

        for start_idx in field_starts:
            field_name = None
            separate_idx = None
            end_idx = None
            for j in range(start_idx + 1, min(start_idx + 100, len(all_elems))):
                elem = all_elems[j]
                if elem.tag == qn('w:instrText'):
                    match = re.search(r'MERGEFIELD\s+"?([^"\\]+)"?', elem.text or '')
                    if match:
                        field_name = match.group(1).strip()
                elif elem.tag == qn('w:fldChar'):
                    fld_type = elem.get(qn('w:fldCharType'))
                    if fld_type == 'separate':
                        separate_idx = j
                    elif fld_type == 'end':
                        end_idx = j
                        break

            if field_name and separate_idx and end_idx:
                value = field_map.get(field_name)
                if value is None:
                    continue
                first_set = False
                for j in range(separate_idx + 1, end_idx):
                    if all_elems[j].tag == qn('w:t'):
                        if not first_set:
                            all_elems[j].text = str(value)
                            first_set = True
                        else:
                            all_elems[j].text = ''

        for elem in list(element.iter()):
            if elem.tag == qn('w:fldSimple'):
                instr = elem.get(qn('w:instr'), '')
                match = re.search(r'MERGEFIELD\s+"?([^"\\]+)"?', instr)
                if match:
                    fname = match.group(1).strip()
                    value = field_map.get(fname)
                    if value is not None:
                        for run in elem.findall(qn('w:r')):
                            for t in run.findall(qn('w:t')):
                                t.text = str(value)

    process_element(doc.element.body)
    for section in doc.sections:
        for attr in ['header', 'footer', 'first_page_header', 'first_page_footer']:
            try:
                part = getattr(section, attr, None)
                if part and part._element is not None:
                    process_element(part._element)
            except Exception:
                pass


CHART_POSITIONS = [43, 84, 127, 160, 189, 212, 236]
SUMMARY_CHART_POSITION = 269


def replace_chart_at_paragraph(doc, para_index, chart_path):
    if para_index >= len(doc.paragraphs):
        return
    para = doc.paragraphs[para_index]
    for run in para.runs:
        for drawing in run._element.findall(qn('w:drawing')):
            run._element.remove(drawing)
    for child in list(para._element):
        if 'drawing' in child.tag or 'object' in child.tag.lower():
            para._element.remove(child)
    for run in para.runs:
        run.text = ''
    run = para.runs[0] if para.runs else para.add_run()
    run.add_picture(chart_path, width=Cm(10))


def generate_single_report(entity_name, backend_bytes, template_bytes, work_dir):
    """Generate report for one entity, return (word_bytes, entity_name)."""
    # Recalc
    work_file = recalc_for_entity(backend_bytes, entity_name, work_dir)

    # Read data
    fields = read_all_fields(work_file)
    blocks, summary_pcts = read_chart_data(work_file)

    # Generate charts
    chart_dir = os.path.join(work_dir, "charts")
    os.makedirs(chart_dir, exist_ok=True)
    chart_paths = []
    for i, block in enumerate(blocks):
        path = os.path.join(chart_dir, f"bloque_{i+1}.png")
        generate_block_chart(block, entity_name.strip(), path)
        chart_paths.append(path)
    summary_path = os.path.join(chart_dir, "resumen.png")
    generate_summary_chart(summary_pcts, entity_name.strip(), summary_path)

    # Fill template
    doc = Document(BytesIO(template_bytes))
    fill_merge_fields(doc, fields)

    # Insert charts
    for i, (para_idx, chart_path) in enumerate(zip(CHART_POSITIONS, chart_paths)):
        try:
            replace_chart_at_paragraph(doc, para_idx, chart_path)
        except Exception:
            pass
    try:
        replace_chart_at_paragraph(doc, SUMMARY_CHART_POSITION, summary_path)
    except Exception:
        pass

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Streamlit UI ───

st.title("📊 Generador de Informes de Madurez Digital")
st.markdown("Sube el Excel backend y la plantilla Word, selecciona las entidades y genera todos los informes automáticamente.")

st.divider()

# File uploads
col1, col2 = st.columns(2)
with col1:
    st.subheader("1️⃣ Excel Backend")
    backend_file = st.file_uploader(
        "Sube el archivo Excel con los datos del formulario",
        type=['xlsx'],
        key='backend'
    )

with col2:
    st.subheader("2️⃣ Plantilla Word")
    template_file = st.file_uploader(
        "Sube la plantilla del informe (.docx)",
        type=['docx'],
        key='template'
    )

st.divider()

if backend_file and template_file:
    backend_bytes = backend_file.read()
    template_bytes = template_file.read()

    # Read entities
    entities = get_entities(backend_bytes)

    st.subheader("3️⃣ Selecciona las entidades")

    col_sel1, col_sel2 = st.columns([3, 1])
    with col_sel2:
        if st.button("✅ Seleccionar todas"):
            st.session_state['selected'] = list(range(len(entities)))
        if st.button("❌ Deseleccionar todas"):
            st.session_state['selected'] = []

    with col_sel1:
        selected_entities = st.multiselect(
            "Elige las entidades para las que generar informes:",
            options=entities,
            default=entities,
            format_func=lambda x: x.strip()
        )

    st.divider()

    if selected_entities:
        st.subheader(f"4️⃣ Generar {len(selected_entities)} informes")

        if st.button(f"🚀 Generar {len(selected_entities)} informes", type="primary", use_container_width=True):

            progress = st.progress(0, text="Preparando...")
            results = []
            errors = []

            with tempfile.TemporaryDirectory() as work_dir:
                for i, entity in enumerate(selected_entities):
                    progress.progress(
                        (i) / len(selected_entities),
                        text=f"Procesando {entity.strip()} ({i+1}/{len(selected_entities)})..."
                    )
                    try:
                        word_bytes = generate_single_report(
                            entity, backend_bytes, template_bytes, work_dir
                        )
                        safe_name = re.sub(r'[^\w\s-]', '', entity.strip()).strip().replace(' ', '_')[:50]
                        safe_name = safe_name.encode('ascii', 'ignore').decode('ascii') or f"entidad_{i}"
                        results.append((f"Informe_MD_{safe_name}.docx", word_bytes))
                    except Exception as e:
                        errors.append((entity.strip(), str(e)))

            progress.progress(1.0, text="✅ ¡Completado!")

            # Show results
            st.success(f"✅ {len(results)} informes generados correctamente")

            if errors:
                st.warning(f"⚠️ {len(errors)} informes con errores:")
                for ent, err in errors:
                    st.text(f"  • {ent}: {err}")

            # Create ZIP
            if results:
                zip_buf = BytesIO()
                with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for filename, content in results:
                        zf.writestr(filename, content)
                zip_buf.seek(0)

                st.download_button(
                    label=f"📥 Descargar {len(results)} informes (ZIP)",
                    data=zip_buf,
                    file_name=f"Informes_Madurez_Digital_{datetime.now().strftime('%Y%m%d')}.zip",
                    mime="application/zip",
                    use_container_width=True
                )

                # Also offer individual downloads
                with st.expander("📄 Descargar informes individuales"):
                    for filename, content in results:
                        st.download_button(
                            label=f"📄 {filename}",
                            data=content,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=filename
                        )
    else:
        st.info("Selecciona al menos una entidad para generar informes.")

else:
    st.info("👆 Sube los dos archivos para empezar: el Excel backend y la plantilla Word del informe.")

# Footer
st.divider()
st.caption("Generador de Informes de Madurez Digital — DCA en RED | Fundación Esplai")
